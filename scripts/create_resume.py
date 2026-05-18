#!/usr/bin/env python3
"""Generate a tailored resume DOCX from a base template.

The LLM agent (Claude CLI) calls this with the per-job summary, title line, and
selected certifications. This script:

  1. Validates the summary against the CV WRITING RULES — rejects any text
     containing logistics phrases (sponsorship, availability, time-zone,
     engagement-platform, remote-mode, bilingual, application-flow,
     location-as-availability). Defense in depth: even if the LLM forgets
     the prompt rules, the DOCX never gets written with bad content.
  2. Loads the candidate's profile from config/profile.yaml.
  3. Opens the base resume template at $BASE_DIR/resume_base.docx.
  4. Replaces the title-line and summary paragraphs.
  5. Updates the Certifications section to match the selected IDs.
  6. Writes the tailored DOCX to the output path.

Usage:
  python scripts/create_resume.py \
      --output   "$BASE_DIR/Jobs/Title - Company/Resume - You - Company.docx" \
      --job-type "power"  \
      --title-line "Power Platform Developer  |  Power Automate · Python" \
      --summary  "Automation engineer with 2+ years building Power Automate ..." \
      --certs    "py-crash,pbi-dax,git-atlassian,..."

If --summary fails validation, the script exits 2 with a diff-style error
showing which forbidden phrases were detected and which category each fell
into. The caller (the LLM agent) is expected to fix and retry.
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path

# ──────────────────────────────────────────────────────────────────────
# CV WRITING RULES — canonical forbidden-phrase patterns
# ──────────────────────────────────────────────────────────────────────
# The Professional Summary describes what the candidate IS — role, stack,
# measurable wins. Everything else (logistics) belongs in OTHER sections:
#   • Languages       → languages section of the resume
#   • Availability    → screening Q&A
#   • Time zone       → contact block
#   • Visa status     → screening Q&A or cover line
#   • Engagement      → never name the platform / agency
#   • Remote mode     → contact block / implicit from job mode
#   • Application     → cover letter content, not summary
#
# Source of truth for these patterns is the cv-writing-rules / cv-summary-
# cleaner skills (installed in Claude Desktop). When a new forbidden phrase
# is added to those skills, add it here too — both must stay in sync.

FORBIDDEN_CATEGORIES: dict[str, list[str]] = {
    "sponsorship": [
        r"\bno\s+(visa\s+)?sponsorship\s+(required|needed)\b",
        r"\bdoes\s+not\s+require\s+sponsorship\b",
        r"\bsin\s+(necesidad\s+de\s+)?patrocinio\b",
    ],
    "availability": [
        r"\bavailable\s+immediately\b",
        r"\bimmediately\s+available\b",
        r"\bdisponibilidad\s+inmediata\b",
        r"\bdisponible\s+(de\s+)?inmediat[ao]\b",
    ],
    "time_zone_alignment": [
        r"\bUS[-\s]?(business[-\s]?hours|timezone)\s+aligned\b",
        r"\b(eastern|EST|EDT|PST|PDT|CST|CDT|GMT)[-\s]?\d?\s*(overlap|aligned)\b",
        r"\b(time[-\s]?zone)\s+(aligned|overlap)\b",
    ],
    "engagement_platform": [
        r"\bvia\s+(toptal|revelo|torre|hirelatam|fiverr|upwork)\b",
        r"\bseeking\s+(freelance|nearshore|remote)\s+engagements?\s+via\b",
        r"\bavailable\s+for\s+(immediate\s+)?placement\s+via\b",
    ],
    "remote_mode_credential": [
        r"\b100\s?%\s+remote\b",
        r"\bfully\s+remote\s+(from|with|ready)\b",
        r"\bremote[-\s]?first\s+ready\b",
        r"\bthrive\s+in\s+(fully\s+)?remote\s+environments?\b",
        r"\b(listo|preparado)\s+(para\s+integrar(se|me))\s+de\s+forma\s+remota\b",
    ],
    "bilingual": [
        r"\bbilingual\s+(english|spanish)[\s/-]*(english|spanish)?\s*\(?[Cc]2\)?\b",
        r"\bbilingüe\s+(español|inglés)\b",
        r"\binglés\s+C2\s+y\s+español\s+nativo\b",
        r"\bprofesional\s+bilingüe\b",
    ],
    "application_flow": [
        r"\bI'?m\s+applying\s+for\b",
        r"\bI\s+am\s+applying\s+(for|to)\b",
        r"\baplico\s+a\s+este\s+rol\b",
        r"\bposiciono\s+esta\s+candidatura\b",
        r"\bbusco\s+ahora\s+aportar\b",
    ],
    "location_as_availability": [
        # Pairs "based in X" or "X-based" with availability/timezone phrasing.
        # Both patterns are location-agnostic — they match the logistics
        # *structure*, not a specific country, so the rule generalizes to
        # any candidate using this repo.
        r"\bbased\s+in\s+\w+[\s,]+(seeking|with|US[-\s]?(timezone|business)|available)\b",
        r"\b\w+[-\s]?based[-\s,]+(seeking|US[-\s]?timezone)\b",
    ],
}

COMPILED = {
    cat: [re.compile(p, re.IGNORECASE) for p in patterns]
    for cat, patterns in FORBIDDEN_CATEGORIES.items()
}


def validate_summary(summary: str) -> list[tuple[str, str]]:
    """Scan `summary` against every forbidden category and return a list of
    (category, offending_phrase) tuples. Empty list means the summary passes.

    This is intentionally a regex-only check — no NLP, no LLM. If the LLM
    rephrases cleverly enough to slip past these patterns, that's fine; the
    rules are about the *category* of content, and the regex catches the
    common phrasings. Add new patterns to FORBIDDEN_CATEGORIES as needed."""
    hits: list[tuple[str, str]] = []
    for category, patterns in COMPILED.items():
        for pat in patterns:
            m = pat.search(summary)
            if m:
                hits.append((category, m.group(0)))
    return hits


def _require_env(name: str) -> str:
    val = os.environ.get(name)
    if not val:
        raise SystemExit(
            f"{name} not set — source your .env before running this script."
        )
    return val


# ──────────────────────────────────────────────────────────────────────
# DOCX manipulation — minimal scaffold
# ──────────────────────────────────────────────────────────────────────
# The example here uses python-docx to clone a base template and replace
# the title-line and summary paragraphs. Production users typically extend
# this with: full styling, certification table, achievement bullets, etc.
# Treat this as the "skeleton" — the value-add is the validation gate
# above, which prevents bad summaries from ever reaching a DOCX.


def write_resume(output_path: str, base_template: str,
                 title_line: str, summary: str,
                 certs: list[str]) -> None:
    """Open `base_template`, replace the title-line and summary paragraphs,
    update the certifications section, and save to `output_path`.

    Customize the paragraph-detection logic to match your template's
    structure. The default scans for paragraphs marked with stable text
    tokens — `{{TITLE_LINE}}` and `{{SUMMARY}}` — which the template
    placeholder approach is the most maintainable pattern."""
    try:
        from docx import Document
    except ImportError:
        raise SystemExit(
            "python-docx not installed. Run: pip install -r requirements.txt"
        ) from None

    if not os.path.exists(base_template):
        raise SystemExit(f"Base resume template not found: {base_template}")

    doc = Document(base_template)

    for para in doc.paragraphs:
        if "{{TITLE_LINE}}" in para.text:
            _replace_paragraph_text(para, title_line)
        elif "{{SUMMARY}}" in para.text:
            _replace_paragraph_text(para, summary)
        elif "{{CERTS}}" in para.text:
            _replace_paragraph_text(para, "  ·  ".join(certs))

    Path(os.path.dirname(output_path)).mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def _replace_paragraph_text(para, new_text: str) -> None:
    """Replace the entire text of a paragraph with `new_text`, preserving
    the formatting of the first run. python-docx doesn't have a one-liner
    for this; we clear all runs then add a single new run."""
    if not para.runs:
        para.add_run(new_text)
        return
    first_run = para.runs[0]
    first_run.text = new_text
    for run in para.runs[1:]:
        run.text = ""


# ──────────────────────────────────────────────────────────────────────
# CLI entry
# ──────────────────────────────────────────────────────────────────────


def main() -> int:
    ap = argparse.ArgumentParser(
        description="Generate a tailored resume DOCX, enforcing CV writing rules."
    )
    ap.add_argument("--output", required=True, help="Path to write the tailored DOCX")
    ap.add_argument("--job-type", required=True,
                    choices=["n8n", "power", "ai", "general"],
                    help="Variant of the resume to generate")
    ap.add_argument("--title-line", required=True, help="Resume header title line")
    ap.add_argument("--summary", required=True, help="Professional summary text")
    ap.add_argument("--certs", default="",
                    help="Comma-separated certification IDs")
    args = ap.parse_args()

    # ─── Step 1: validate the summary BEFORE doing any I/O ───
    violations = validate_summary(args.summary)
    if violations:
        print("\n❌ Summary failed CV WRITING RULES validation.\n", file=sys.stderr)
        print("The following forbidden phrases were found:", file=sys.stderr)
        for category, phrase in violations:
            print(f"  • [{category}]  '{phrase}'", file=sys.stderr)
        print("\nRewrite the summary to remove these phrases.", file=sys.stderr)
        print("Reference: docs/architecture.md → CV writing rules.", file=sys.stderr)
        return 2

    # ─── Step 2: resolve paths from env ───
    base_dir = _require_env("BASE_DIR")
    base_template = os.path.join(base_dir, "resume_base.docx")

    certs = [c.strip() for c in args.certs.split(",") if c.strip()]
    write_resume(args.output, base_template, args.title_line, args.summary, certs)

    print(f"Wrote {args.output} ({args.job_type}, {len(certs)} certs)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
